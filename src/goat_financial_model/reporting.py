"""Shared export helpers for goat farming workbooks and reports."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, Optional, Sequence

import numpy as np
import pandas as pd


DEFAULT_MODEL_AUTHOR = "Goat Farmers United"


@dataclass
class ChartSpec:
    title: str
    kind: str
    data: pd.DataFrame
    y_axis_title: str = ""
    stacked: bool = False


@dataclass
class ExportBundle:
    model: Any
    scenario_name: str
    author_name: str
    base_df: pd.DataFrame = field(default_factory=pd.DataFrame)
    scenario_df: pd.DataFrame = field(default_factory=pd.DataFrame)
    annual_df: pd.DataFrame = field(default_factory=pd.DataFrame)
    kpis_df: pd.DataFrame = field(default_factory=pd.DataFrame)
    break_even_df: pd.DataFrame = field(default_factory=pd.DataFrame)
    statements: Dict[str, pd.DataFrame] = field(default_factory=dict)
    statement_errors: Dict[str, str] = field(default_factory=dict)
    supplementary: Dict[str, pd.DataFrame] = field(default_factory=dict)
    scenario_inputs: Dict[str, Any] = field(default_factory=dict)
    export_timestamp_utc: datetime = field(
        default_factory=lambda: datetime.now(timezone.utc)
    )


def _clean_author_name(author_name: Optional[str]) -> str:
    value = str(author_name or "").strip()
    return value or DEFAULT_MODEL_AUTHOR


def _ensure_dataframe(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame()
    return df.copy()


def _coerce_datetime_index(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    work = df.copy()
    if isinstance(work.index, pd.DatetimeIndex):
        return work
    if "Period" in work.columns:
        periods = pd.to_datetime(work["Period"], errors="coerce")
        if periods.notna().any():
            work = work.drop(columns=["Period"])
            work.index = pd.DatetimeIndex(periods)
            work.index.name = "Period"
            return work
    return work


def _annualize_schedule(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    work = _coerce_datetime_index(df)
    if not isinstance(work.index, pd.DatetimeIndex):
        return pd.DataFrame()
    numeric = work.apply(pd.to_numeric, errors="coerce")
    annual = numeric.groupby(work.index.year).sum(min_count=1)
    annual.index.name = "Year"
    return annual


def _build_statements(
    model: Any,
    scenario_df: pd.DataFrame,
) -> tuple[Dict[str, pd.DataFrame], Dict[str, str]]:
    statements: Dict[str, pd.DataFrame] = {}
    errors: Dict[str, str] = {}
    if scenario_df.empty:
        return statements, errors
    for title, builder_name in (
        ("Statement of Financial Performance", "statement_of_financial_performance"),
        ("Statement of Financial Position", "statement_of_financial_position"),
        ("Statement of Cash Flows", "statement_of_cash_flow"),
    ):
        try:
            builder = getattr(model, builder_name)
            frame = builder(scenario_df, annual=True)
        except (AttributeError, KeyError, TypeError, ValueError) as exc:
            errors[title] = str(exc) or "Statement build failed."
            continue
        if isinstance(frame, pd.DataFrame) and not frame.empty:
            statements[title] = frame.copy()
    return statements, errors


def prepare_export_bundle(
    model: Any,
    *,
    scenario_name: str,
    author_name: Optional[str] = None,
    base_df: Optional[pd.DataFrame] = None,
    scenario_df: Optional[pd.DataFrame] = None,
    kpis_df: Optional[pd.DataFrame] = None,
    break_even_df: Optional[pd.DataFrame] = None,
    supplementary: Optional[Dict[str, pd.DataFrame]] = None,
    scenario_inputs: Optional[Dict[str, Any]] = None,
) -> ExportBundle:
    base_frame = _ensure_dataframe(base_df)
    scenario_frame = _ensure_dataframe(scenario_df)
    annual_df = _annualize_schedule(scenario_frame)
    kpis_frame = _ensure_dataframe(kpis_df)
    if kpis_frame.empty and not scenario_frame.empty:
        try:
            kpis_frame = _ensure_dataframe(model.kpis(scenario_frame, annual=True))
        except ValueError:
            kpis_frame = pd.DataFrame()
    break_even_frame = _ensure_dataframe(break_even_df)
    if break_even_frame.empty and not scenario_frame.empty:
        try:
            break_even_frame = _ensure_dataframe(model.break_even(scenario_frame, annual=True))
        except ValueError:
            break_even_frame = pd.DataFrame()
    statements, statement_errors = _build_statements(model, scenario_frame)
    return ExportBundle(
        model=model,
        scenario_name=scenario_name,
        author_name=_clean_author_name(author_name),
        base_df=base_frame,
        scenario_df=scenario_frame,
        annual_df=annual_df,
        kpis_df=kpis_frame,
        break_even_df=break_even_frame,
        statements=statements,
        statement_errors=statement_errors,
        supplementary={
            str(name): table.copy()
            for name, table in (supplementary or {}).items()
            if isinstance(table, pd.DataFrame) and not table.empty
        },
        scenario_inputs=dict(scenario_inputs or {}),
    )


def _prepare_dataframe_for_excel(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    frame = _ensure_dataframe(df)
    if frame.empty:
        return frame
    if isinstance(frame.columns, pd.MultiIndex):
        frame.columns = [
            " - ".join(str(part) for part in entry if str(part))
            for entry in frame.columns.to_flat_index()
        ]
    if isinstance(frame.index, pd.MultiIndex):
        frame = frame.reset_index()
    elif isinstance(frame.index, pd.DatetimeIndex):
        frame = frame.reset_index()
        first_col = frame.columns[0]
        frame = frame.rename(columns={first_col: "Period"})
        frame["Period"] = pd.to_datetime(frame["Period"], errors="coerce")
    elif not isinstance(frame.index, pd.RangeIndex) or frame.index.name:
        frame = frame.reset_index()
        first_col = frame.columns[0]
        if first_col == "index":
            frame = frame.rename(columns={first_col: str(frame.index.name or "Index")})
    return frame


def _sanitize_sheet_name(name: str, existing: set[str]) -> str:
    invalid = set('[]:*?/\\')
    cleaned = "".join("_" if ch in invalid else ch for ch in name).strip() or "Sheet"
    cleaned = cleaned[:31]
    base = cleaned
    suffix = 1
    while cleaned in existing:
        candidate = f"{base[:31 - len(str(suffix)) - 1]}_{suffix}"
        cleaned = candidate[:31]
        suffix += 1
    existing.add(cleaned)
    return cleaned


def _latest_value(series: Optional[pd.Series]) -> Optional[Any]:
    if series is None:
        return None
    clean = pd.to_numeric(series, errors="coerce")
    clean = clean.dropna()
    if clean.empty:
        return None
    return clean.iloc[-1]


def _summary_rows(bundle: ExportBundle) -> list[tuple[str, Any, str]]:
    rows: list[tuple[str, Any, str]] = [
        ("Scenario", bundle.scenario_name, "text"),
        (
            "Prepared On (UTC)",
            bundle.export_timestamp_utc.strftime("%Y-%m-%d %H:%M"),
            "text",
        ),
        ("Prepared By", bundle.author_name, "text"),
    ]
    annual = bundle.annual_df
    if not annual.empty:
        for label, column in (
            ("Latest Revenue", "Revenue"),
            ("Latest EBITDA", "EBITDA"),
            ("Latest Net Profit", "NPAT"),
        ):
            if column in annual.columns:
                value = _latest_value(annual[column])
                if value is not None:
                    rows.append((label, float(value), "currency"))
    kpis = bundle.kpis_df
    if not kpis.empty:
        latest = kpis.iloc[-1]
        for label, kind in (
            ("Gross Margin %", "percent"),
            ("EBITDA Margin %", "percent"),
            ("Net Margin %", "percent"),
            ("NPV", "currency"),
            ("IRR", "percent"),
            ("DSCR", "ratio"),
            ("Interest Coverage", "ratio"),
        ):
            if label in latest.index and pd.notna(latest[label]):
                rows.append((label, latest[label], kind))
    if not bundle.break_even_df.empty and "Break-even Revenue" in bundle.break_even_df.columns:
        value = _latest_value(bundle.break_even_df["Break-even Revenue"])
        if value is not None:
            rows.append(("Break-even Revenue", float(value), "currency"))
    return rows[:12]


def _metadata_frame(bundle: ExportBundle) -> pd.DataFrame:
    horizon = len(bundle.scenario_df.index) if not bundle.scenario_df.empty else 0
    start_period = ""
    end_period = ""
    if not bundle.scenario_df.empty:
        work = _coerce_datetime_index(bundle.scenario_df)
        if isinstance(work.index, pd.DatetimeIndex) and not work.index.empty:
            start_period = work.index.min().strftime("%Y-%m-%d")
            end_period = work.index.max().strftime("%Y-%m-%d")
    rows = [
        ("Scenario", bundle.scenario_name),
        ("Prepared By", bundle.author_name),
        ("Export Timestamp (UTC)", bundle.export_timestamp_utc.strftime("%Y-%m-%d %H:%M:%S")),
        ("Timeline Periods", horizon),
        ("Start Period", start_period),
        ("End Period", end_period),
    ]
    for key, value in bundle.scenario_inputs.items():
        rows.append((str(key), value))
    return pd.DataFrame(rows, columns=["Field", "Value"])


def _statement_warnings_frame(bundle: ExportBundle) -> pd.DataFrame:
    if not bundle.statement_errors:
        return pd.DataFrame()
    return pd.DataFrame(
        [
            {"Statement": statement, "Error": error}
            for statement, error in bundle.statement_errors.items()
        ]
    )


def _comparison_chart_frame(
    base_df: pd.DataFrame,
    scenario_df: pd.DataFrame,
    columns: Sequence[str],
) -> pd.DataFrame:
    base_annual = _annualize_schedule(base_df)
    scenario_annual = _annualize_schedule(scenario_df)
    frames: list[pd.DataFrame] = []
    if not base_annual.empty:
        base = base_annual[[col for col in columns if col in base_annual.columns]].copy()
        base.columns = [f"Base {col}" for col in base.columns]
        frames.append(base)
    if not scenario_annual.empty:
        scenario = scenario_annual[[col for col in columns if col in scenario_annual.columns]].copy()
        scenario.columns = [f"Scenario {col}" for col in scenario.columns]
        frames.append(scenario)
    if not frames:
        return pd.DataFrame()
    chart = pd.concat(frames, axis=1).sort_index()
    chart.index.name = "Year"
    return chart


def _margin_chart_frame(kpis_df: pd.DataFrame) -> pd.DataFrame:
    if kpis_df.empty:
        return pd.DataFrame()
    columns = [
        col
        for col in ["Gross Margin %", "EBITDA Margin %", "Net Margin %", "IRR"]
        if col in kpis_df.columns
    ]
    if not columns:
        return pd.DataFrame()
    chart = kpis_df[columns].apply(pd.to_numeric, errors="coerce") * 100.0
    chart.index.name = str(kpis_df.index.name or "Year")
    return chart


def _cash_flow_chart_frame(annual_df: pd.DataFrame) -> pd.DataFrame:
    if annual_df.empty:
        return pd.DataFrame()
    columns = [
        col for col in ["CFO", "CFI", "CFF", "Net Cash Flow"] if col in annual_df.columns
    ]
    if not columns:
        return pd.DataFrame()
    chart = annual_df[columns].apply(pd.to_numeric, errors="coerce")
    chart.index.name = str(annual_df.index.name or "Year")
    return chart


def _balance_chart_frame(annual_df: pd.DataFrame) -> pd.DataFrame:
    if annual_df.empty:
        return pd.DataFrame()
    chart = pd.DataFrame(index=annual_df.index)
    current_assets = pd.to_numeric(annual_df.get("Current Assets"), errors="coerce")
    non_current_assets = pd.to_numeric(annual_df.get("Non-current Assets"), errors="coerce")
    current_liabilities = pd.to_numeric(annual_df.get("Current Liabilities"), errors="coerce")
    non_current_liabilities = pd.to_numeric(
        annual_df.get("Non-current Liabilities"), errors="coerce"
    )
    equity = pd.to_numeric(annual_df.get("Equity"), errors="coerce")
    if current_assets.notna().any() or non_current_assets.notna().any():
        chart["Total Assets"] = current_assets.fillna(0.0) + non_current_assets.fillna(0.0)
    if current_liabilities.notna().any() or non_current_liabilities.notna().any():
        chart["Total Liabilities"] = current_liabilities.fillna(0.0) + non_current_liabilities.fillna(0.0)
    if equity.notna().any():
        chart["Equity"] = equity
    chart = chart.dropna(axis=1, how="all")
    chart.index.name = str(annual_df.index.name or "Year")
    return chart


def _break_even_chart_frame(break_even_df: pd.DataFrame) -> pd.DataFrame:
    if break_even_df.empty:
        return pd.DataFrame()
    columns = [
        col
        for col in ["Break-even Revenue", "Fixed Costs (approx)"]
        if col in break_even_df.columns
    ]
    if not columns:
        return pd.DataFrame()
    chart = break_even_df[columns].apply(pd.to_numeric, errors="coerce")
    chart.index.name = str(break_even_df.index.name or "Year")
    return chart


def build_chart_specs(bundle: ExportBundle) -> list[ChartSpec]:
    specs: list[ChartSpec] = []
    comparison = _comparison_chart_frame(
        bundle.base_df,
        bundle.scenario_df,
        ("Revenue", "EBITDA", "NPAT"),
    )
    if not comparison.empty:
        specs.append(
            ChartSpec(
                title="Base vs Scenario Profit Trend",
                kind="line",
                data=comparison,
                y_axis_title="Value",
            )
        )
    margins = _margin_chart_frame(bundle.kpis_df)
    if not margins.empty:
        specs.append(
            ChartSpec(
                title="Margin and IRR Profile",
                kind="line",
                data=margins,
                y_axis_title="Percent",
            )
        )
    cash_flows = _cash_flow_chart_frame(bundle.annual_df)
    if not cash_flows.empty:
        specs.append(
            ChartSpec(
                title="Cash Flow Bridge",
                kind="bar",
                data=cash_flows,
                y_axis_title="Cash flow",
            )
        )
    balance = _balance_chart_frame(bundle.annual_df)
    if not balance.empty:
        specs.append(
            ChartSpec(
                title="Balance Sheet Overview",
                kind="line",
                data=balance,
                y_axis_title="Balance",
            )
        )
    break_even = _break_even_chart_frame(bundle.break_even_df)
    if not break_even.empty:
        specs.append(
            ChartSpec(
                title="Break-even Revenue vs Fixed Costs",
                kind="bar",
                data=break_even,
                y_axis_title="Value",
            )
        )
    return specs


def _header_text(value: Any) -> str:
    return str(value or "").strip()


def _classify_column_format(header: str) -> Optional[str]:
    lower = header.lower()
    if lower in {"period", "date"} or lower.endswith(" period"):
        return "yyyy-mm-dd"
    if lower == "year":
        return "0"
    if (
        "%" in header
        or "margin" in lower
        or "growth" in lower
        or "allocation" in lower
        or lower == "irr"
    ):
        return "0.0%"
    if lower in {"dscr", "interest coverage"}:
        return '0.00"x"'
    if "payback" in lower:
        return "0.0"
    if any(
        token in lower
        for token in (
            "revenue",
            "cogs",
            "expense",
            "cost",
            "profit",
            "ebit",
            "cash",
            "asset",
            "liabilit",
            "equity",
            "capex",
            "npv",
            "value",
            "salary",
            "price",
            "reserve",
        )
    ):
        return '#,##0.00;[Red]-#,##0.00'
    if any(token in lower for token in ("heads", "months", "quantity", "units")):
        return '#,##0.00'
    return '#,##0.00'


def _style_overview_sheet(ws, accent: str, accent_soft: str, bundle: ExportBundle, sheet_name_map: Dict[str, str]) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill

    ws["A1"] = "Goat Farm Financial Model"
    ws["A1"].font = Font(size=20, bold=True, color="1F2937")
    ws["A2"] = (
        "Scenario workbook with linked tabs for inputs, annual performance, "
        "financial statements, charts, and export metadata."
    )
    ws["A2"].font = Font(size=11, color="4B5563")
    ws["A4"] = "Executive Snapshot"
    ws["A4"].font = Font(size=12, bold=True, color=accent)
    ws["A5"] = "Metric"
    ws["B5"] = "Value"
    for cell in ws[5]:
        cell.fill = PatternFill("solid", fgColor=accent)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row_idx, (label, value, kind) in enumerate(_summary_rows(bundle), start=6):
        ws.cell(row=row_idx, column=1, value=label)
        value_cell = ws.cell(row=row_idx, column=2, value=value)
        if kind == "currency":
            value_cell.number_format = '#,##0.00;[Red]-#,##0.00'
        elif kind == "percent":
            value_cell.number_format = "0.0%"
        elif kind == "ratio":
            value_cell.number_format = '0.00"x"'
    ws["D4"] = "Workbook Guide"
    ws["D4"].font = Font(size=12, bold=True, color=accent)
    guide_rows = [
        ("Charts", "View the packaged chart set and source tables."),
        ("Metadata", "Review scenario assumptions and export controls."),
    ]
    for offset, (label, description) in enumerate(guide_rows, start=5):
        ws.cell(row=offset, column=4, value=label)
        ws.cell(row=offset, column=5, value=description)
    ws["D9"] = "Jump to Sheets"
    ws["D9"].font = Font(size=12, bold=True, color=accent)
    row_idx = 10
    for logical_name, actual_name in sheet_name_map.items():
        if actual_name == ws.title:
            continue
        cell = ws.cell(row=row_idx, column=4, value=logical_name)
        cell.hyperlink = f"#'{actual_name}'!A1"
        cell.style = "Hyperlink"
        ws.cell(row=row_idx, column=5, value=actual_name)
        row_idx += 1
    ws.freeze_panes = "A6"
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["D"].width = 28
    ws.column_dimensions["E"].width = 44
    for row_idx in range(6, min(ws.max_row, 18) + 1):
        if row_idx % 2 == 0:
            for cell in ws[row_idx][:2]:
                cell.fill = PatternFill("solid", fgColor=accent_soft)


def _style_data_sheet(ws, accent: str, accent_soft: str) -> None:
    from datetime import date, datetime as dt_datetime

    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    ws.sheet_view.showGridLines = False
    if ws.max_row > 1:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.fill = PatternFill("solid", fgColor=accent)
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.auto_filter.ref = ws.dimensions
        for row_idx in range(2, min(ws.max_row, 200) + 1):
            if row_idx % 2 == 0:
                for cell in ws[row_idx]:
                    cell.fill = PatternFill("solid", fgColor=accent_soft)
    for col_idx in range(1, ws.max_column + 1):
        header = _header_text(ws.cell(row=1, column=col_idx).value)
        number_format = _classify_column_format(header)
        max_length = len(header)
        for row_idx in range(2, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            value = cell.value
            if value is None:
                continue
            if isinstance(value, (pd.Timestamp, dt_datetime, date)):
                cell.number_format = "yyyy-mm-dd"
            elif isinstance(value, (int, float, np.integer, np.floating)) and number_format:
                cell.number_format = number_format
            max_length = max(max_length, len(str(value)))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_length + 2, 12), 32)


def _write_dataframe_to_sheet(ws, df: pd.DataFrame, start_row: int, start_col: int = 1) -> tuple[int, int]:
    frame = _prepare_dataframe_for_excel(df)
    if frame.empty and frame.columns.empty:
        return start_row, start_col
    for col_offset, column in enumerate(frame.columns, start=start_col):
        ws.cell(row=start_row, column=col_offset, value=str(column))
    for row_offset, (_, row) in enumerate(frame.iterrows(), start=start_row + 1):
        for col_offset, value in enumerate(row.tolist(), start=start_col):
            ws.cell(row=row_offset, column=col_offset, value=value)
    return start_row, start_col + len(frame.columns) - 1


def _add_excel_chart(ws, spec: ChartSpec, start_row: int) -> int:
    from openpyxl.chart import BarChart, LineChart, Reference

    if spec.data.empty:
        return start_row
    source = spec.data.reset_index()
    source.columns = [str(col) for col in source.columns]
    ws.cell(row=start_row, column=1, value=spec.title)
    ws.cell(row=start_row, column=1).style = "Title"
    header_row, end_col = _write_dataframe_to_sheet(ws, source, start_row + 1)
    data_start_row = start_row + 2
    data_end_row = start_row + 1 + len(source)
    if data_end_row < data_start_row:
        return start_row + 4
    if spec.kind == "bar":
        chart = BarChart()
        chart.type = "col"
        chart.grouping = "clustered"
        chart.overlap = 0
    else:
        chart = LineChart()
    chart.title = spec.title
    chart.style = 2
    chart.height = 7
    chart.width = 11
    if spec.y_axis_title:
        chart.y_axis.title = spec.y_axis_title
    chart.x_axis.title = source.columns[0]
    data_ref = Reference(
        ws,
        min_col=2,
        max_col=end_col,
        min_row=header_row,
        max_row=data_end_row,
    )
    cats_ref = Reference(
        ws,
        min_col=1,
        max_col=1,
        min_row=data_start_row,
        max_row=data_end_row,
    )
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats_ref)
    ws.add_chart(chart, f"G{start_row}")
    return start_row + max(len(source) + 4, 18)


def generate_excel_workbook(bundle: ExportBundle) -> bytes:
    from openpyxl.styles import Font

    buffer = BytesIO()
    used_sheets: set[str] = set()
    sheet_name_map: Dict[str, str] = {}
    accent = "0F766E"
    accent_soft = "E8F7F3"
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        workbook = writer.book
        workbook.properties.creator = bundle.author_name
        workbook.properties.title = "Goat Farm Financial Model"
        workbook.properties.subject = bundle.scenario_name

        def write_sheet(logical_name: str, df: pd.DataFrame) -> Optional[str]:
            frame = _prepare_dataframe_for_excel(df)
            if frame.empty and frame.columns.empty:
                return None
            actual_name = _sanitize_sheet_name(logical_name, used_sheets)
            frame.to_excel(writer, sheet_name=actual_name, index=False)
            sheet_name_map[logical_name] = actual_name
            return actual_name

        write_sheet("Input Schedule", bundle.base_df)
        write_sheet(f"{bundle.scenario_name} Timeline", bundle.scenario_df)
        write_sheet(f"{bundle.scenario_name} Annual", bundle.annual_df)
        write_sheet("KPIs (Annual)", bundle.kpis_df)
        write_sheet("Break-even Analysis", bundle.break_even_df)
        write_sheet("Scenario Inputs", pd.DataFrame(
            {
                "Input": list(bundle.scenario_inputs.keys()),
                "Value": list(bundle.scenario_inputs.values()),
            }
        ))
        write_sheet("Model Metadata", _metadata_frame(bundle))
        write_sheet("Statement Warnings", _statement_warnings_frame(bundle))
        for statement_name, statement_df in bundle.statements.items():
            write_sheet(statement_name, statement_df)
        for name, table in bundle.supplementary.items():
            write_sheet(f"Supplementary - {name}", table)

        overview_name = _sanitize_sheet_name("Overview", used_sheets)
        charts_name = _sanitize_sheet_name("Charts", used_sheets)
        overview = workbook.create_sheet(overview_name, 0)
        charts_sheet = workbook.create_sheet(charts_name, 1)
        sheet_name_map["Overview"] = overview_name
        sheet_name_map["Charts"] = charts_name
        _style_overview_sheet(overview, accent, accent_soft, bundle, sheet_name_map)

        charts_sheet["A1"] = "Chart Pack"
        charts_sheet["A1"].font = Font(size=18, bold=True, color="1F2937")
        charts_sheet["A2"] = "Each chart includes its underlying source table for auditability."
        current_row = 4
        for spec in build_chart_specs(bundle):
            current_row = _add_excel_chart(charts_sheet, spec, current_row)

        for worksheet in workbook.worksheets:
            if worksheet.title == overview_name:
                continue
            _style_data_sheet(worksheet, accent, accent_soft)

    return buffer.getvalue()


def _matplotlib_chart_image(spec: ChartSpec) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    if spec.data.empty:
        return b""
    fig, ax = plt.subplots(figsize=(8.5, 4.2))
    data = spec.data.copy()
    if spec.kind == "bar":
        data.plot(kind="bar", ax=ax)
    else:
        data.plot(ax=ax, marker="o")
    ax.set_title(spec.title)
    ax.set_xlabel(str(data.index.name or "Period"))
    if spec.y_axis_title:
        ax.set_ylabel(spec.y_axis_title)
    ax.grid(axis="y", alpha=0.25)
    ax.ticklabel_format(style="plain", axis="y")
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    output = BytesIO()
    fig.savefig(output, format="png", dpi=180)
    plt.close(fig)
    return output.getvalue()


def build_pdf_chart_flowables(bundle: ExportBundle) -> list[Any]:
    from reportlab.lib.units import cm
    from reportlab.platypus import Image, Spacer

    flowables: list[Any] = []
    for spec in build_chart_specs(bundle):
        image_bytes = _matplotlib_chart_image(spec)
        if not image_bytes:
            continue
        flowables.append(Image(BytesIO(image_bytes), width=16 * cm, height=8 * cm))
        flowables.append(Spacer(1, 10))
    return flowables


def build_docx_report(
    bundle: ExportBundle,
    *,
    title: str,
    subtitle: str,
) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    document = Document()
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = document.styles["Title"].font.size

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run(subtitle)

    document.add_heading("Executive Summary", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Light List Accent 1"
    summary_table.rows[0].cells[0].text = "Metric"
    summary_table.rows[0].cells[1].text = "Value"
    for label, value, kind in _summary_rows(bundle):
        row = summary_table.add_row().cells
        row[0].text = label
        if kind == "currency" and isinstance(value, (int, float, np.integer, np.floating)):
            row[1].text = f"{float(value):,.2f}"
        elif kind == "percent" and isinstance(value, (int, float, np.integer, np.floating)):
            row[1].text = f"{float(value):.1%}"
        elif kind == "ratio" and isinstance(value, (int, float, np.integer, np.floating)):
            row[1].text = f"{float(value):.2f}x"
        else:
            row[1].text = str(value)

    document.add_heading("Scenario Charts", level=1)
    for spec in build_chart_specs(bundle):
        document.add_heading(spec.title, level=2)
        image_bytes = _matplotlib_chart_image(spec)
        if image_bytes:
            document.add_picture(BytesIO(image_bytes), width=Inches(6.5))

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Key Tables", level=1)
    for title_text, frame in (
        ("Scenario Timeline", bundle.scenario_df.head(12)),
        ("Annual KPIs", bundle.kpis_df.head(12)),
        ("Break-even Analysis", bundle.break_even_df.head(12)),
    ):
        if frame.empty:
            continue
        document.add_heading(title_text, level=2)
        export_frame = _prepare_dataframe_for_excel(frame)
        table = document.add_table(rows=1, cols=len(export_frame.columns))
        table.style = "Light List Accent 1"
        for idx, column in enumerate(export_frame.columns):
            table.rows[0].cells[idx].text = str(column)
        for _, row in export_frame.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row.tolist()):
                cells[idx].text = "" if pd.isna(value) else str(value)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
